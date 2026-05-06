"""Production-grade DOCX report generator for compliance scan results.

Generates a professionally styled Word document with:
- Cover page with branding
- Color-coded risk levels (red/orange/yellow/green)
- Styled table headers with dark backgrounds
- Executive summary callout boxes
- Posture score visualization with color coding
- Landscape orientation for gap tables
- Proper typography and spacing
"""

from __future__ import annotations

from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Color palette
_NAVY = RGBColor(0x23, 0x2F, 0x3E)
_DARK_BLUE = RGBColor(0x1A, 0x47, 0x7A)
_LIGHT_BLUE = RGBColor(0xD6, 0xEA, 0xF8)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_RED = RGBColor(0xE7, 0x4C, 0x3C)
_ORANGE = RGBColor(0xF3, 0x9C, 0x12)
_YELLOW = RGBColor(0xF1, 0xC4, 0x0F)
_GREEN = RGBColor(0x27, 0xAE, 0x60)
_GRAY = RGBColor(0x7F, 0x8C, 0x8D)
_LIGHT_GRAY = RGBColor(0xEC, 0xF0, 0xF1)
_DARK_GRAY = RGBColor(0x2C, 0x3E, 0x50)

_RISK_COLORS = {
    "critical": "E74C3C",
    "high": "E67E22",
    "medium": "F39C12",
    "low": "3498DB",
}

_CONFIDENCE_COLORS = {
    "high": "27AE60",
    "medium": "F39C12",
    "low": "E74C3C",
}

_POSTURE_COLORS = {
    "pass": "27AE60",
    "warn": "F39C12",
    "fail": "E74C3C",
}


def _shade_cell(cell, color_hex: str) -> None:
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 8,
                   color: RGBColor | None = None, align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                      risk_col: int | None = None, confidence_col: int | None = None,
                      small: bool = False) -> None:
    """Add a professionally styled table with colored header and optional risk/confidence coloring."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Style header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade_cell(hdr_cells[i], "1A477A")
        _set_cell_text(hdr_cells[i], h, bold=True, size=8 if small else 9, color=_WHITE,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Data rows with alternating shading
    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            # Alternating row background
            if row_idx % 2 == 0:
                _shade_cell(row_cells[i], "F8F9FA")

            # Risk column coloring
            if i == risk_col and val.lower() in _RISK_COLORS:
                _shade_cell(row_cells[i], _RISK_COLORS[val.lower()])
                _set_cell_text(row_cells[i], val.upper(), bold=True, size=7 if small else 8,
                               color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            # Confidence column coloring
            elif i == confidence_col and val.lower() in _CONFIDENCE_COLORS:
                _shade_cell(row_cells[i], _CONFIDENCE_COLORS[val.lower()])
                _set_cell_text(row_cells[i], val.upper(), bold=True, size=7 if small else 8,
                               color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _set_cell_text(row_cells[i], val, size=7 if small else 8)

        row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # Remove table borders and add subtle grid
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
        f'<w:insideH w:val="single" w:sz="2" w:space="0" w:color="D5D8DC"/>'
        f'<w:insideV w:val="single" w:sz="2" w:space="0" w:color="D5D8DC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()


def _add_posture_table(doc: Document, posture_data: list[tuple[str, dict]]) -> None:
    """Add a posture score table with color-coded scores."""
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    headers = ["Framework", "Score", "Satisfied", "Total", "Status"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade_cell(hdr_cells[i], "1A477A")
        _set_cell_text(hdr_cells[i], h, bold=True, size=9, color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for label, posture in posture_data:
        if not posture:
            continue
        score = posture.get("score", 0)
        satisfied = posture.get("satisfied", posture.get("covered_domains", 0))
        total = posture.get("total", 0)

        if score >= 80:
            status = "PASS"
            color_hex = "27AE60"
        elif score >= 50:
            status = "NEEDS WORK"
            color_hex = "F39C12"
        else:
            status = "CRITICAL"
            color_hex = "E74C3C"

        row_cells = table.add_row().cells
        _set_cell_text(row_cells[0], label, bold=True, size=9)
        _shade_cell(row_cells[1], color_hex)
        _set_cell_text(row_cells[1], f"{score}%", bold=True, size=10, color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row_cells[2], str(satisfied), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row_cells[3], str(total), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(row_cells[4], color_hex)
        _set_cell_text(row_cells[4], status, bold=True, size=8, color=_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()


def _add_cover_page(doc: Document, data: dict) -> None:
    """Add a professional cover page."""
    # Add spacing
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AWS India Regulatory\nCompliance Assessment Report")
    run.font.size = Pt(28)
    run.font.color.rgb = _DARK_BLUE
    run.bold = True

    doc.add_paragraph()

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("DPDP Act 2023  |  RBI Master Direction  |  SEBI CSCRF  |  CERT-In Directions")
    run.font.size = Pt(12)
    run.font.color.rgb = _GRAY

    for _ in range(3):
        doc.add_paragraph()

    # Metadata box
    meta = data.get("scan_metadata", {})
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Region: ").bold = True
    info.add_run(f"{data.get('region', 'us-east-1')}\n")
    info.add_run(f"Scope: ").bold = True
    info.add_run(f"{data.get('aggregator', 'single-account')}\n")
    info.add_run(f"Scan Date: ").bold = True
    info.add_run(f"{meta.get('scan_start', datetime.utcnow().isoformat())[:10]}\n")
    info.add_run(f"Resources Scanned: ").bold = True
    info.add_run(f"{data.get('total_components', 0)}\n")
    info.add_run(f"Compliance Gaps: ").bold = True
    info.add_run(f"{data.get('total_gaps', 0)}")

    for _ in range(4):
        doc.add_paragraph()

    # Disclaimer
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run("CONFIDENTIAL - For internal compliance review only")
    run.font.size = Pt(9)
    run.font.color.rgb = _RED
    run.italic = True

    doc.add_page_break()


def _add_section_header(doc: Document, text: str) -> None:
    """Add a styled section header with underline."""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = _DARK_BLUE


def _set_landscape(doc: Document) -> None:
    """Set all sections to landscape."""
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1.5)


def generate_docx(data: dict[str, Any], ct_data: dict[str, Any] | None = None) -> Document:
    """Generate a production-grade DOCX compliance report.

    Args:
        data: Account/org scan result dict.
        ct_data: Optional Control Tower scan result dict.

    Returns:
        A styled python-docx Document object.
    """
    doc = Document()
    _set_landscape(doc)

    # Cover page
    _add_cover_page(doc, data)

    # Executive Summary
    _add_section_header(doc, "Executive Summary")
    p = doc.add_paragraph()
    run = p.add_run(data.get("executive_summary", ""))
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Key metrics callout
    p = doc.add_paragraph()
    p.add_run(f"Total Resources: ").bold = True
    p.add_run(f"{data.get('total_components', 0)}     ")
    p.add_run(f"Total Gaps: ").bold = True
    p.add_run(f"{data.get('total_gaps', 0)}     ")
    p.add_run(f"Suppressed: ").bold = True
    p.add_run(f"{data.get('suppressed_count', 0)}     ")
    conf = data.get("confidence_distribution", {})
    if conf:
        p.add_run(f"High Confidence: ").bold = True
        p.add_run(f"{conf.get('high', 0)}     ")
        p.add_run(f"Medium: ").bold = True
        p.add_run(f"{conf.get('medium', 0)}     ")
        p.add_run(f"Low: ").bold = True
        p.add_run(f"{conf.get('low', 0)}")

    doc.add_paragraph()

    # Compliance Posture
    _add_section_header(doc, "Compliance Posture Scores")
    posture_items = [
        ("DPDP Act 2023 (Digital Personal Data Protection)", data.get("dpdp_posture")),
        ("RBI Master Direction (IT Governance)", data.get("rbi_posture")),
        ("SEBI CSCRF 2024 (Cyber Security & Resilience)", data.get("sebi_posture")),
        ("CERT-In Directions 2022 (Incident Reporting)", data.get("certin_posture")),
    ]
    _add_posture_table(doc, posture_items)

    # Confidence Distribution
    if conf:
        _add_section_header(doc, "Confidence Distribution")
        doc.add_paragraph("Each compliance gap is assigned a confidence level indicating the reliability of the finding.")
        _add_styled_table(doc, ["Level", "Count", "Description"], [
            ["HIGH", str(conf.get("high", 0)), "Direct technical check verifiable from AWS Config (e.g., encryption disabled)"],
            ["MEDIUM", str(conf.get("medium", 0)), "Interpretive mapping from regulatory requirement to AWS control (e.g., data localization)"],
            ["LOW", str(conf.get("low", 0)), "Organizational requirement where infrastructure is only a proxy (e.g., DPO appointment)"],
        ], confidence_col=0)

    # Gap Summary by Framework
    gap_summary = data.get("gap_summary_by_framework", {})
    if gap_summary:
        _add_section_header(doc, "Gap Summary by Framework and Risk")

        # Penalty exposure summary
        penalty_dist: dict[str, int] = {}
        for g in data.get("gaps", []):
            p = g.get("penalty_exposure", "")
            if p:
                penalty_dist[p] = penalty_dist.get(p, 0) + 1
        if penalty_dist:
            doc.add_heading("Potential Penalty Exposure", level=2)
            penalty_rows = [[k, str(v)] for k, v in sorted(penalty_dist.items(), key=lambda x: -x[1])]
            _add_styled_table(doc, ["Penalty Category", "Gap Count"], penalty_rows)

        rows = []
        for fw, risks in sorted(gap_summary.items()):
            c = risks.get("critical", 0)
            h = risks.get("high", 0)
            m = risks.get("medium", 0)
            lo = risks.get("low", 0)
            rows.append([fw.upper(), str(c), str(h), str(m), str(lo), str(c+h+m+lo)])
        _add_styled_table(doc, ["Framework", "Critical", "High", "Medium", "Low", "Total"], rows)

    # Resource-Level Compliance
    rc = data.get("resource_compliance", {})
    if rc:
        _add_section_header(doc, "Resource-Level Compliance by Domain")
        doc.add_paragraph("Pass rate indicates the percentage of checked resources that meet the domain requirements.")
        rows = []
        for domain_key in sorted(rc.keys()):
            d = rc[domain_key]
            pct = d.get("pct", 0)
            status = "PASS" if pct >= 80 else "NEEDS WORK" if pct >= 50 else "FAILING"
            rows.append([domain_key, str(d["checked"]), str(d["passed"]), str(d["failed"]), f"{pct}%", status])
        _add_styled_table(doc, ["Domain", "Checked", "Passed", "Failed", "Pass Rate", "Status"], rows)

    # Per-Account Breakdown
    pa = data.get("per_account", {})
    if pa:
        _add_section_header(doc, "Per-Account Compliance Breakdown")
        rows = []
        for acct in sorted(pa.keys()):
            ad = pa[acct]
            dpdp_s = ad.get("dpdp_posture", {}).get("score", "n/a")
            rbi_s = ad.get("rbi_posture", {}).get("score", "n/a")
            sebi_s = ad.get("sebi_posture", {}).get("score", "n/a")
            rows.append([acct, str(ad["gap_count"]), f"{dpdp_s}%", f"{rbi_s}%", f"{sebi_s}%"])
        _add_styled_table(doc, ["Account ID", "Gap Count", "DPDP Score", "RBI Score", "SEBI Score"], rows)

    # Critical Gaps
    gaps = data.get("gaps", [])
    critical_gaps = [g for g in gaps if g.get("risk") == "critical"]
    if critical_gaps:
        doc.add_page_break()
        _add_section_header(doc, f"Critical Gaps ({len(critical_gaps)})")
        doc.add_paragraph("These findings require immediate attention. All are high-confidence direct technical checks.")
        rows = []
        for g in critical_gaps:
            rows.append([
                g.get("component", ""),
                g.get("framework", "").upper(),
                g.get("domain_name", ""),
                g.get("gap", ""),
                g.get("confidence", ""),
                g.get("penalty_exposure", ""),
                g.get("responsibility_type", ""),
            ])
        _add_styled_table(doc, ["Component", "Framework", "Domain", "Finding", "Confidence", "Penalty Exposure", "Responsibility"],
                          rows, confidence_col=4, small=True)

    # High Gaps
    high_gaps = [g for g in gaps if g.get("risk") == "high"]
    if high_gaps:
        doc.add_page_break()
        _add_section_header(doc, f"High-Risk Gaps ({len(high_gaps)} total, showing first 80)")
        rows = []
        for g in high_gaps[:80]:
            rows.append([
                g.get("component", ""),
                g.get("framework", "").upper(),
                g.get("domain_name", ""),
                g.get("gap", ""),
                g.get("confidence", ""),
                g.get("penalty_exposure", ""),
                g.get("responsibility_type", ""),
            ])
        _add_styled_table(doc, ["Component", "Framework", "Domain", "Finding", "Confidence", "Penalty Exposure", "Responsibility"],
                          rows, confidence_col=4, small=True)
        if len(high_gaps) > 80:
            doc.add_paragraph(f"... and {len(high_gaps) - 80} additional high-risk gaps. See full JSON report for complete list.")

    # Medium/Low summary
    medium_gaps = [g for g in gaps if g.get("risk") == "medium"]
    low_gaps = [g for g in gaps if g.get("risk") == "low"]
    if medium_gaps or low_gaps:
        _add_section_header(doc, "Medium & Low Risk Summary")
        if medium_gaps:
            doc.add_paragraph(f"Medium risk: {len(medium_gaps)} gaps (90-180 day remediation window)")
        if low_gaps:
            doc.add_paragraph(f"Low risk: {len(low_gaps)} gaps (advisory, organizational controls)")

    # Suppressed Gaps
    suppressed = data.get("suppressed_gaps", [])
    if suppressed:
        _add_section_header(doc, "Suppressed Gaps (Exceptions)")
        doc.add_paragraph("These gaps were auto-suppressed based on exception rules. Included for audit transparency.")
        rows = []
        for s in suppressed:
            rows.append([s.get("component", ""), s.get("framework", "").upper(),
                         s.get("gap", ""), s.get("suppression_reason", "")])
        _add_styled_table(doc, ["Component", "Framework", "Finding", "Suppression Reason"], rows, small=True)

    # Remediation Timeline
    timeline = data.get("remediation_timeline", [])
    if timeline:
        doc.add_page_break()
        _add_section_header(doc, "Remediation Timeline")
        for phase in timeline:
            h = doc.add_heading(phase["phase"], level=2)
            for run in h.runs:
                run.font.color.rgb = _DARK_GRAY
            for item in phase.get("items", []):
                doc.add_paragraph(item, style="List Bullet")

    # --- Control Tower Section ---
    if ct_data:
        doc.add_page_break()
        title = doc.add_heading("Control Tower Governance Assessment", level=0)
        for run in title.runs:
            run.font.color.rgb = _DARK_BLUE

        doc.add_paragraph(ct_data.get("executive_summary", ""))

        # Landing Zone
        lz = ct_data.get("landing_zone", {})
        if lz:
            _add_section_header(doc, "Landing Zone Status")
            _add_styled_table(doc, ["Property", "Value"], [
                ["Status", lz.get("status", "unknown")],
                ["Version", lz.get("version", "unknown")],
                ["Drift Status", lz.get("drift_status", "unknown")],
                ["Total OUs", str(ct_data.get("total_ous", 0))],
                ["Enabled Controls", str(ct_data.get("total_enabled_controls", 0))],
            ])

        # CT Posture
        _add_section_header(doc, "Control Tower Domain Coverage")
        ct_posture = [
            ("DPDP Act 2023", ct_data.get("dpdp_posture")),
            ("RBI Master Direction", ct_data.get("rbi_posture")),
            ("SEBI CSCRF", ct_data.get("sebi_posture")),
            ("CERT-In Directions", ct_data.get("certin_posture")),
        ]
        _add_posture_table(doc, ct_posture)

        # Per-OU
        per_ou = ct_data.get("per_ou", {})
        if per_ou:
            _add_section_header(doc, "Per-OU Control Coverage")
            rows = []
            for ou_name, od in sorted(per_ou.items()):
                rows.append([ou_name, str(od.get("enabled_controls", 0)),
                             str(od.get("dpdp_domains_covered", 0)), str(od.get("rbi_domains_covered", 0)),
                             str(od.get("sebi_domains_covered", 0)), str(od.get("certin_domains_covered", 0))])
            _add_styled_table(doc, ["OU Name", "Controls Enabled", "DPDP Domains", "RBI Domains", "SEBI Domains", "CERT-In Domains"], rows)

        # CT Gaps with missing controls
        ct_gaps = ct_data.get("gaps", [])
        if ct_gaps:
            _add_section_header(doc, "Missing Guardrails by Framework")
            doc.add_paragraph("These Control Tower guardrails are not currently enabled but are required for compliance coverage.")
            rows = []
            for g in ct_gaps:
                missing = ", ".join(g.get("missing_controls", []))
                rows.append([g.get("framework", "").upper(), g.get("domain_name", ""),
                             g.get("gap", ""), missing, g.get("confidence", "")])
            _add_styled_table(doc, ["Framework", "Domain", "Gap", "Missing Controls", "Confidence"],
                              rows, confidence_col=4, small=True)

        # Recommendations
        recs = ct_data.get("recommendations", [])
        if recs:
            _add_section_header(doc, "Recommended Controls to Enable")
            doc.add_paragraph("Prioritized list of Control Tower guardrails to enable for improved compliance posture.")
            rows = []
            for r in recs:
                rows.append([r.get("priority", "").upper(), r.get("control_id", ""),
                             r.get("description", ""), r.get("framework", "").upper(), r.get("domain_name", "")])
            _add_styled_table(doc, ["Priority", "Control ID", "Description", "Framework", "Domain"], rows)

    # Footer disclaimer
    doc.add_page_break()
    _add_section_header(doc, "Disclaimer")
    disc = doc.add_paragraph()
    disc.add_run(
        "This report provides automated assessment guidance based on published regulatory frameworks. "
        "It does not constitute legal advice or compliance certification. Organizations should consult "
        "qualified compliance and legal professionals for definitive regulatory compliance determinations.\n\n"
        "Confidence levels indicate the reliability of each finding:\n"
        "- HIGH: Verifiable technical fact from AWS Config data\n"
        "- MEDIUM: Interpretive mapping requiring human review\n"
        "- LOW: Organizational requirement that cannot be verified from infrastructure\n\n"
        "Generated by AWS India Compliance MCP Server v0.1.0"
    ).font.size = Pt(9)

    return doc
